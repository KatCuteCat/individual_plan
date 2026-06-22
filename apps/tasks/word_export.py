"""
Генератор Word-документа индивидуального плана

"""

from io import BytesIO
from decimal import Decimal

from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


#  Константы

FONT_NAME = 'Times New Roman'

MONTH_GENITIVE = [
    '', 'января', 'февраля', 'марта', 'апреля', 'мая', 'июня',
    'июля', 'августа', 'сентября', 'октября', 'ноября', 'декабря',
]

CATEGORY_SECTIONS = [
    ('EDUCATIONAL_METHODICAL', '2. Учебно-методическая работа'),
    ('SCIENTIFIC', '3. Научно-исследовательская работа'),
    ('ORGANIZATIONAL', '4. Организационно-методическая работа'),
    ('EDUCATIONAL_WORK', '5. Воспитательная работа'),
]

CENTER = WD_ALIGN_PARAGRAPH.CENTER
RIGHT = WD_ALIGN_PARAGRAPH.RIGHT
LEFT = WD_ALIGN_PARAGRAPH.LEFT

# Основной размер шрифта документа
SZ_MAIN = Pt(14)
# Размер шрифта для сводной таблицы на стр.2
SZ_SUMMARY = Pt(14)
# Размер шрифта для таблицы 1.1 (сводная учебная)
SZ_T11 = Pt(9)
# Размер шрифта для таблицы 1.2 (детализация по дисциплинам)
SZ_T12 = Pt(9)
# Размер шрифта для таблиц разделов 2–5
SZ_TASK = Pt(11)
# Размер шрифта для заголовков разделов 2–5
SZ_SECTION_TITLE = Pt(12)
# Размер шрифта для подписей-подсказок («Подпись», «(должность)» и т.п.)
SZ_HINT = Pt(12)


# Форматирование

def _fmt(value):
    """Десятичное число → строка с запятой: 892.50 → '892,5'; None/0 → '0'."""
    if value is None:
        return '0'
    try:
        d = Decimal(str(value)).quantize(Decimal('0.01'))
    except Exception:
        return str(value)
    if d == 0:
        return '0'
    s = str(d)
    if '.' in s:
        s = s.rstrip('0').rstrip('.')
    return s.replace('.', ',')


def _fmt_date(dt):
    """date → 'дд.мм.гггг' или плейсхолдер."""
    if dt is None:
        return '__.__.____'
    return dt.strftime('%d.%m.%Y')


def _date_signature(dt):
    """date → 'dd.mm.YYYY' или плейсхолдер '__.__.____'."""
    if dt is None:
        return '__.__.____'
    return f'{dt.day:02d}.{dt.month:02d}.{dt.year}'


def _protocol_line(settings, use_report=False):
    """Строка протокола (план или отчёт)."""
    num = '     '
    date_str = '__.__.____'
    if settings:
        if use_report:
            proto_num = getattr(settings, 'report_protocol_number', '') or ''
            proto_date = getattr(settings, 'report_protocol_date', None)
        else:
            proto_num = settings.protocol_number or ''
            proto_date = settings.protocol_date
        if proto_num:
            num = proto_num
        if proto_date:
            date_str = (f'{proto_date.day:02d}.{proto_date.month:02d}.'
                        f'{proto_date.year}')
    return f'протокол № {num} от {date_str} г.'


def _s(settings, field, placeholder='________'):
    """Значение поля ExportSettings или плейсхолдер."""
    if settings is None:
        return placeholder
    val = getattr(settings, field, '')
    return val if val else placeholder


def _get_date(settings, field):
    """Получить дату из ExportSettings или None."""
    if settings is None:
        return None
    return getattr(settings, field, None)




def _set_run_font(run, size=SZ_MAIN, bold=False, underline=False):
    """Установить шрифт Times New Roman на run (включая кириллицу)."""
    run.font.name = FONT_NAME
    run.font.size = size
    run.bold = bold
    run.underline = underline
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rFonts.set(qn(attr), FONT_NAME)


def _para(doc, text='', size=SZ_MAIN, bold=False, underline=False,
          align=None, after=Pt(0), before=Pt(0), left_indent=None):
    """Добавить параграф с форматированием. Полуторный интервал."""
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = after
    pf.space_before = before
    pf.line_spacing = 1.5
    if left_indent is not None:
        pf.left_indent = left_indent
    if text:
        run = p.add_run(text)
        _set_run_font(run, size, bold, underline)
    return p


def _run(para, text, size=SZ_MAIN, bold=False, underline=False):
    """Добавить run к существующему параграфу."""
    run = para.add_run(text)
    _set_run_font(run, size, bold, underline)
    return run


def _cell(cell, text, size=SZ_T11, bold=False, underline=False,
          align=LEFT):
    """Записать текст в ячейку таблицы. Одинарный интервал."""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    if text is not None and str(text):
        run = p.add_run(str(text))
        _set_run_font(run, size, bold, underline)


def _cell_two_lines(cell, line1, line2, size=SZ_TASK, align=CENTER):
    """Две строки в одной ячейке (через soft break)."""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r1 = p.add_run(str(line1))
    _set_run_font(r1, size)
    r1.add_break()
    r2 = p.add_run(str(line2))
    _set_run_font(r2, size)


def _set_widths(table, widths):
    """Установить ширину столбцов."""
    table.autofit = False
    for row in table.rows:
        for idx, w in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = w


def _remove_all_borders(table):
    """Убрать все границы у таблицы (невидимая таблица)."""
    tbl = table._element
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        elem = OxmlElement(f'w:{edge}')
        elem.set(qn('w:val'), 'none')
        elem.set(qn('w:sz'), '0')
        elem.set(qn('w:space'), '0')
        elem.set(qn('w:color'), 'auto')
        borders.append(elem)
    tblPr.append(borders)


def _set_cell_bottom_border(cell, sz='4', color='000000'):
    """Установить только нижнюю границу ячейки (для линий подписи)."""
    tc = cell._element
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    tcBorders = OxmlElement('w:tcBorders')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), sz)
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), color)
    tcBorders.append(bottom)
    tcPr.append(tcBorders)


# Настройка документа

def _setup(doc):
    """Формат страницы A4, поля, шрифт по умолчанию."""
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Mm(15)
    sec.bottom_margin = Mm(15)
    sec.left_margin = Mm(25)
    sec.right_margin = Mm(10)
    # Ширина контента = 210 − 25 − 10 = 175 мм

    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = SZ_MAIN
    pf = style.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.5


# Страница 1: Шапка

def _add_header_page(doc, pos, settings):
    # Шапка университета — 14pt, по центру
    _para(doc,
          'Министерство науки и высшего образования Российской Федерации',
          SZ_MAIN, align=CENTER)
    _para(doc,
          'Федеральное государственное бюджетное образовательное учреждение '
          'высшего образования',
          SZ_MAIN, align=CENTER)
    _para(doc,
          '«Воронежский государственный технический университет»',
          SZ_MAIN, align=CENTER)

    # 2 пустых строки после шапки
    _para(doc)
    _para(doc)

    approve_table = doc.add_table(rows=6, cols=2)
    _remove_all_borders(approve_table)
    _set_widths(approve_table, [Mm(100), Mm(75)])

    # Строка 0: УТВЕРЖДАЮ (жирный, по левому краю линий)
    _cell(approve_table.cell(0, 0), '', SZ_MAIN)
    _cell(approve_table.cell(0, 1), 'УТВЕРЖДАЮ',
          SZ_MAIN, bold=True, align=LEFT)

    # Строка 1: Должность утверждающего + нижняя граница ячейки
    approver_title = _s(settings, 'approver_title', '')
    title_text = approver_title if approver_title != '________' else ''
    _cell(approve_table.cell(1, 0), '', SZ_MAIN)
    _cell(approve_table.cell(1, 1), title_text, SZ_MAIN, align=LEFT)
    _set_cell_bottom_border(approve_table.cell(1, 1))

    # Строка 2: подсказка «(должность)»
    _cell(approve_table.cell(2, 0), '', SZ_MAIN)
    _cell(approve_table.cell(2, 1), '(должность)',
          SZ_HINT, align=CENTER)

    # Строка 3: ФИО утверждающего + нижняя граница ячейки
    approver_name = _s(settings, 'approver_name', '')
    name_text = approver_name if approver_name != '________' else ''
    _cell(approve_table.cell(3, 0), '', SZ_MAIN)
    _cell(approve_table.cell(3, 1), name_text, SZ_MAIN, align=LEFT)
    _set_cell_bottom_border(approve_table.cell(3, 1))

    # Строка 4: подсказка «(фамилия, имя, отчество)»
    _cell(approve_table.cell(4, 0), '', SZ_MAIN)
    _cell(approve_table.cell(4, 1), '(фамилия, имя, отчество)',
          SZ_HINT, align=CENTER)

    # Строка 5: дата (из protocol_date или плейсхолдер)
    plan_date = _get_date(settings, 'protocol_date')
    _cell(approve_table.cell(5, 0), '', SZ_MAIN)
    _cell(approve_table.cell(5, 1),
          f'{_date_signature(plan_date)} г.',
          SZ_MAIN, align=LEFT)

    # 2 пустых строки перед «ИНДИВИДУАЛЬНЫЙ ПЛАН»
    _para(doc)
    _para(doc)

    # Заголовок плана — жирный, по центру
    _para(doc, 'ИНДИВИДУАЛЬНЫЙ ПЛАН',
          SZ_MAIN, bold=True, align=CENTER)
    _para(doc, 'РАБОТЫ ПРЕПОДАВАТЕЛЯ',
          SZ_MAIN, bold=True, align=CENTER)

    # 1 пустая строка
    _para(doc)

    # Кафедра — по левому краю, название подчёркнуто
    dept = _s(settings, 'department_full_name', '___________')
    p = _para(doc, 'Кафедра ', SZ_MAIN)
    _run(p, dept, SZ_MAIN, underline=True)

    # Учебный год — по центру
    year_name = str(pos.academic_year.name)
    _para(doc, f'на {year_name} учебный год',
          SZ_MAIN, align=CENTER)

    # 1 пустая строка
    _para(doc)

    # ФИО — значение подчёркнуто
    full_name = pos.user.full_name or pos.user.get_full_name() or '________'
    p = _para(doc, 'Фамилия, имя, отчество ', SZ_MAIN)
    _run(p, full_name, SZ_MAIN, underline=True)

    # Должность + вид занятости — значение подчёркнуто
    pos_name = str(pos.position.name) if pos.position else '________'
    empl = pos.get_employment_type_display()
    p = _para(doc, 'Должность ', SZ_MAIN)
    _run(p, f'{pos_name} {empl}', SZ_MAIN, underline=True)

    # Учёное звание, степень — 14pt; значение рядом без подчёркивания
    acad = pos.user.academic_title or ''
    if acad:
        _para(doc, f'Учёное звание, степень {acad}', SZ_MAIN)
    else:
        _para(doc, 'Учёное звание, степень', SZ_MAIN)

    doc.add_page_break()


# Страница 2: Сводная таблица

def _add_summary_page(doc, pos, settings, stats, plan_total, fact_total):
    # Заголовок — единственный жирный текст на странице
    _para(doc, 'СВОДНЫЕ ДАННЫЕ ПО ПЛАНУ',
          SZ_MAIN, bold=True, align=CENTER)


    from django.db.models import Sum
    from apps.tasks.models import Task

    cat_rows = (
        Task.objects.filter(
            assignee=pos,
            status__in=[Task.STATUS_APPROVED, Task.STATUS_COMPLETED],
        )
        .values('category__code')
        .annotate(plan=Sum('planned_hours'), fact=Sum('actual_hours'))
    )
    cat_map = {
        row['category__code']: (
            row['plan'] or Decimal('0'),
            row['fact'] or Decimal('0'),
        )
        for row in cat_rows
    }


    rows_data = [('1. Учебная работа', plan_total, fact_total)]
    if not pos.is_hourly:
        for code, label in CATEGORY_SECTIONS:
            p, f = cat_map.get(code, (Decimal('0'), Decimal('0')))
            rows_data.append((label, p, f))

    grand_plan = sum(r[1] for r in rows_data)
    grand_fact = sum(r[2] for r in rows_data)

    # ── Таблица: 14pt, без жирного, обычные границы ──
    n = 2 + len(rows_data) + 1
    table = doc.add_table(rows=n, cols=3)
    table.style = 'Table Grid'
    _set_widths(table, [Mm(120), Mm(25), Mm(25)])

    # Заголовок строка 0
    _cell(table.cell(0, 0), 'Объем нагрузки по видам деятельности',
          SZ_SUMMARY, align=LEFT)
    table.cell(0, 1).merge(table.cell(0, 2))
    _cell(table.cell(0, 1), 'За год, часов',
          SZ_SUMMARY, align=CENTER)

    # Заголовок строка 1
    _cell(table.cell(1, 0), '', SZ_SUMMARY)
    _cell(table.cell(1, 1), 'план', SZ_SUMMARY, align=CENTER)
    _cell(table.cell(1, 2), 'факт', SZ_SUMMARY, align=CENTER)

    # Строки данных
    for i, (name, plan, fact) in enumerate(rows_data):
        r = i + 2
        _cell(table.cell(r, 0), name, SZ_SUMMARY)
        _cell(table.cell(r, 1), _fmt(plan), SZ_SUMMARY, align=CENTER)
        _cell(table.cell(r, 2), _fmt(fact), SZ_SUMMARY, align=CENTER)

    # Всего — без жирного
    tr = len(rows_data) + 2
    _cell(table.cell(tr, 0), 'Всего:', SZ_SUMMARY)
    _cell(table.cell(tr, 1), _fmt(grand_plan), SZ_SUMMARY, align=CENTER)
    _cell(table.cell(tr, 2), _fmt(grand_fact), SZ_SUMMARY, align=CENTER)

    # 3 пустых строки после таблицы
    _para(doc)
    _para(doc)
    _para(doc)

    # Протокол (план) — обычный шрифт
    _para(doc, 'План рассмотрен и одобрен на заседании кафедры',
          SZ_MAIN)

    # Строка протокола — дата подчёркнута
    plan_date = _get_date(settings, 'protocol_date')
    p = _para(doc, 'протокол № ', SZ_MAIN)
    proto_num = ''
    if settings:
        proto_num = settings.protocol_number or ''
    _run(p, proto_num if proto_num else '     ', SZ_MAIN)
    _run(p, ' от ', SZ_MAIN)
    _run(p, f'{_date_signature(plan_date)} г.', SZ_MAIN)

    # 1 пустая строка
    _para(doc)

    # Заведующий кафедрой + линия для подписи + дата
    # Без ФИО на этой странице
    p = _para(doc, 'Заведующий кафедрой ____________________',
              SZ_MAIN)
    _run(p, f'{_date_signature(plan_date)} г.', SZ_MAIN)

    # «Подпись» — 12pt, под линией подписи
    _para(doc, 'Подпись', SZ_HINT, left_indent=Mm(55))

    # 1 пустая строка
    _para(doc)

    # Экземпляр плана получил — обычный (не жирный, не подчёркнутый)
    p = _para(doc,
              f'Экземпляр плана получил ________________'
              f'{_date_signature(plan_date)} г.',
              SZ_MAIN)

    # «Подпись преподавателя» — 12pt, под линией
    _para(doc, 'Подпись преподавателя', SZ_HINT, left_indent=Mm(55))

    doc.add_page_break()


#Страница 3: Раздел 1 — Учебная работа

def _add_teaching_section(doc, pos, plan_summary, fact_summary):
    # Заголовок — жирный, подчёркнутый, по центру
    _para(doc, '1. Учебная работа',
          SZ_MAIN, bold=True, underline=True, align=CENTER)

    # 1.1 Сводные данные — жирный, подчёркнутый, по центру
    _para(doc, '1.1 Сводные данные',
          SZ_MAIN, bold=True, underline=True, align=CENTER)

    # Логика данных без изменений
    plan_dict = {r['name']: r['hours'] for r in plan_summary.get('rows', [])}
    fact_dict = {r['name']: r['hours'] for r in fact_summary.get('rows', [])}
    sort_dict = {}
    for r in plan_summary.get('rows', []):
        sort_dict[r['name']] = r['sort_order']
    for r in fact_summary.get('rows', []):
        sort_dict.setdefault(r['name'], r['sort_order'])

    all_names = sorted(
        set(plan_dict) | set(fact_dict),
        key=lambda name: sort_dict.get(name, 999),
    )

    # Таблица 1.1: 9pt, без жирного
    n = 2 + len(all_names) + 1  # 2 заголовка + данные + итого
    t11 = doc.add_table(rows=n, cols=4)
    t11.style = 'Table Grid'
    _set_widths(t11, [Mm(10), Mm(110), Mm(25), Mm(25)])

    # Заголовки — объединяем по вертикали (строки 0-1) для колонок 0 и 1
    t11.cell(0, 0).merge(t11.cell(1, 0))
    _cell(t11.cell(0, 0), '№ п.п', SZ_T11, align=CENTER)

    t11.cell(0, 1).merge(t11.cell(1, 1))
    _cell(t11.cell(0, 1), 'Вид учебной нагрузки',
          SZ_T11, align=CENTER)

    # «Количество часов» объединяем горизонтально
    t11.cell(0, 2).merge(t11.cell(0, 3))
    _cell(t11.cell(0, 2), 'Количество часов',
          SZ_T11, align=CENTER)
    _cell(t11.cell(1, 2), 'план', SZ_T11, align=CENTER)
    _cell(t11.cell(1, 3), 'факт', SZ_T11, align=CENTER)

    # Данные — без жирного
    for i, name in enumerate(all_names):
        r = i + 2
        _cell(t11.cell(r, 0), str(i + 1), SZ_T11, align=CENTER)
        _cell(t11.cell(r, 1), name, SZ_T11)
        _cell(t11.cell(r, 2), _fmt(plan_dict.get(name, 0)),
              SZ_T11, align=CENTER)
        _cell(t11.cell(r, 3), _fmt(fact_dict.get(name, 0)),
              SZ_T11, align=CENTER)

    # Итого — без жирного
    last = n - 1
    _cell(t11.cell(last, 0), '', SZ_T11)
    _cell(t11.cell(last, 1), 'Итого', SZ_T11)
    _cell(t11.cell(last, 2), _fmt(plan_summary.get('total', 0)),
          SZ_T11, align=CENTER)
    _cell(t11.cell(last, 3), _fmt(fact_summary.get('total', 0)),
          SZ_T11, align=CENTER)

    # 1.2 Занятия по учебным дисциплинам
    _para(doc, '1.2 Занятия по учебным дисциплинам',
          SZ_MAIN, bold=True, underline=True, align=CENTER)

    items = list(
        pos.teaching_load
        .select_related('activity_type')
        .order_by('row_num', 'semester', 'discipline',
                  'activity_type__sort_order')
    )

    # Таблица 1.2: всё 9pt, без жирного
    n = 3 + len(items)
    t12 = doc.add_table(rows=n, cols=9)
    t12.style = 'Table Grid'
    _set_widths(t12, [Mm(7), Mm(33), Mm(21), Mm(25),
                      Mm(16), Mm(35), Mm(9), Mm(12), Mm(12)])

    # Строка 0: основные заголовки (cols 0-6 вертикально на 2 строки)
    for col in range(7):
        t12.cell(0, col).merge(t12.cell(1, col))

    # cols 7-8 горизонтально → «Количество часов»
    t12.cell(0, 7).merge(t12.cell(0, 8))

    headers = [
        '№', 'Дисциплина', 'Период контроля',
        'Вид учебной нагрузки', 'Цикл',
        'Факультет, форма обучения, курс, группа',
        'Число студ.',
    ]
    for i, h in enumerate(headers):
        _cell(t12.cell(0, i), h, SZ_T12, align=CENTER)
    _cell(t12.cell(0, 7), 'Количество часов',
          SZ_T12, align=CENTER)

    # Строка 1: план / факт
    _cell(t12.cell(1, 7), 'план', SZ_T12, align=CENTER)
    _cell(t12.cell(1, 8), 'факт', SZ_T12, align=CENTER)

    # Строка 2: номера столбцов 1–9
    for i in range(9):
        _cell(t12.cell(2, i), str(i + 1), SZ_T12, align=CENTER)

    # Данные — 9pt, без жирного
    for idx, item in enumerate(items):
        r = idx + 3
        _cell(t12.cell(r, 0), str(idx + 1), SZ_T12, align=CENTER)
        _cell(t12.cell(r, 1), item.discipline, SZ_T12)
        _cell(t12.cell(r, 2),
              item.period_label or f'{item.semester} сем.', SZ_T12)
        _cell(t12.cell(r, 3),
              item.activity_type.name if item.activity_type else '', SZ_T12)
        _cell(t12.cell(r, 4), item.cycle or '', SZ_T12)
        _cell(t12.cell(r, 5), item.group_number or '', SZ_T12, align=CENTER)
        _cell(t12.cell(r, 6),
              str(item.students_count) if item.students_count else '',
              SZ_T12, align=CENTER)
        _cell(t12.cell(r, 7), _fmt(item.hours), SZ_T12, align=CENTER)
        _cell(t12.cell(r, 8), _fmt(item.hours_fact), SZ_T12, align=CENTER)

    doc.add_page_break()


# Страницы: Разделы 2–5 (задачи)

def _add_task_sections(doc, pos, stats):
    """Разделы 2–5: таблицы задач по категориям. Пропускается для почасовой."""
    from apps.tasks.models import Task

    for section_idx, (cat_code, section_title) in enumerate(CATEGORY_SECTIONS):
        # Заголовок раздела — 12pt, жирный, подчёркнутый, по центру
        _para(doc, section_title,
              SZ_SECTION_TITLE, bold=True, underline=True, align=CENTER,
              before=Pt(10) if section_idx > 0 else Pt(0))

        # В индивидуальный план попадают только реально учтённые задачи:
        # approved (подтверждена зав.) и completed (отчитался, на проверке).
        # assigned / in_progress / pending_approval / declined исключаются.
        tasks = list(
            Task.objects.filter(
                assignee=pos,
                category__code=cat_code,
                status__in=[Task.STATUS_APPROVED, Task.STATUS_COMPLETED],
            ).select_related('work_type').order_by('start_date', 'pk')
        )

        # Таблица: 2 заголовка + 2 строки на задачу (мин. 1) + Всего
        data_count = max(len(tasks), 1)
        n = 2 + (data_count * 2) + 1
        table = doc.add_table(rows=n, cols=6)
        table.style = 'Table Grid'
        _set_widths(table, [Mm(8), Mm(40), Mm(26), Mm(55), Mm(20), Mm(21)])

        # Объединяем заголовочные ячейки вертикально (кроме столбца дат)
        for col in (0, 1, 3, 4, 5):
            table.cell(0, col).merge(table.cell(1, col))

        # Заголовок строка 0 — 11pt, жирный
        h_texts = ['№ п/п', 'ВИДЫ РАБОТ', 'Дата начала',
                   'ОПИСАНИЕ', 'План', 'Факт']
        for i, txt in enumerate(h_texts):
            _cell(table.cell(0, i), txt,
                  SZ_TASK, bold=True, align=CENTER)

        # Заголовок строка 1 — только дата окончания
        _cell(table.cell(1, 2), 'Дата окончания',
              SZ_TASK, bold=True, align=CENTER)

        # Данные — 11pt, жирный
        total_plan = Decimal('0')
        total_fact = Decimal('0')

        if tasks:
            for idx, task in enumerate(tasks):
                r_start = (idx * 2) + 2
                r_end = r_start + 1

                # Объединяем ячейки вертикально (кроме столбца дат)
                for col in (0, 1, 3, 4, 5):
                    table.cell(r_start, col).merge(table.cell(r_end, col))

                # Основная строка данных
                _cell(table.cell(r_start, 0), str(idx + 1),
                      SZ_TASK, bold=True, align=CENTER)

                wt_name = (task.work_type.name
                           if task.work_type else task.title)
                _cell(table.cell(r_start, 1), wt_name, SZ_TASK, bold=True)

                _cell(table.cell(r_start, 2), _fmt_date(task.start_date),
                      SZ_TASK, bold=True, align=CENTER)

                _cell(table.cell(r_start, 3), task.result or '',
                      SZ_TASK, bold=True)

                ph = task.planned_hours or Decimal('0')
                ah = task.actual_hours or Decimal('0')
                _cell(table.cell(r_start, 4), _fmt(ph),
                      SZ_TASK, bold=True, align=CENTER)
                _cell(table.cell(r_start, 5), _fmt(ah),
                      SZ_TASK, bold=True, align=CENTER)

                # Вторая строка — только дата окончания
                _cell(table.cell(r_end, 2), _fmt_date(task.end_date),
                      SZ_TASK, bold=True, align=CENTER)

                total_plan += ph
                total_fact += ah
        else:
            # Пустая строка (2 ряда)
            for col in (0, 1, 3, 4, 5):
                table.cell(2, col).merge(table.cell(3, col))
            for col in range(6):
                _cell(table.cell(2, col), '', SZ_TASK)

        # Всего по разделу — жирный
        last = n - 1
        _cell(table.cell(last, 0), '', SZ_TASK)
        _cell(table.cell(last, 1), 'Всего по разделу',
              SZ_TASK, bold=True)
        _cell(table.cell(last, 2), '', SZ_TASK)
        _cell(table.cell(last, 3), '', SZ_TASK)
        _cell(table.cell(last, 4), _fmt(total_plan),
              SZ_TASK, bold=True, align=CENTER)
        _cell(table.cell(last, 5), _fmt(total_fact),
              SZ_TASK, bold=True, align=CENTER)

    doc.add_page_break()


# Последняя страница: Отчётность

def _add_footer_page(doc, settings):
    # Всё без жирного шрифта на этой странице

    _para(doc, 'ОТЧЁТНОСТЬ ПО ПЛАНУ',
          SZ_MAIN, align=CENTER)

    # 1 пустая строка
    _para(doc)

    _para(doc, 'ОТЧЁТ ПРЕПОДАВАТЕЛЯ',
          SZ_MAIN, align=CENTER)

    # 6 пустых строк
    for _ in range(6):
        _para(doc)

    # Подпись преподавателя — дата из report_protocol_date
    report_date = _get_date(settings, 'report_protocol_date')
    _para(doc,
          f'Преподаватель __________________________'
          f'{_date_signature(report_date)} г.',
          SZ_MAIN)

    # «Подпись» — 12pt, под линией подписи (не под словом «Преподаватель»)
    _para(doc, 'Подпись', SZ_HINT, left_indent=Mm(55))

    # Отчет заслушан
    _para(doc, 'Отчет заслушан на заседании кафедры', SZ_MAIN)
    _para(doc,
          f'({_protocol_line(settings, use_report=True)}).',
          SZ_MAIN)

    # 2 пустых строки
    _para(doc)
    _para(doc)

    # ЗАКЛЮЧЕНИЕ КАФЕДРЫ — без жирного
    _para(doc,
          'ЗАКЛЮЧЕНИЕ КАФЕДРЫ О ВЫПОЛНЕНИИ '
          'ИНДИВИДУАЛЬНОГО ПЛАНА РАБОТЫ ПРЕПОДАВАТЕЛЯ ЗА ГОД',
          SZ_MAIN, align=CENTER)

    # 7 пустых строк
    for _ in range(7):
        _para(doc)

    # Заведующий кафедрой — через невидимую таблицу
    dept_name = _s(settings, 'department_name', '')
    head = _s(settings, 'head_name', '')
    report_date = _get_date(settings, 'report_protocol_date')

    foot_table = doc.add_table(rows=2, cols=3)
    _remove_all_borders(foot_table)
    _set_widths(foot_table, [Mm(53), Mm(83), Mm(38)])

    # Строка 0: «Заведующий кафедрой» | [линия для подписи] | дата
    _cell(foot_table.cell(0, 0), 'Заведующий кафедрой', SZ_MAIN)

    # Среднее поле: аббревиатура + линия + ФИО (или пусто) + нижняя граница
    ds = dept_name if dept_name not in ('', '________', '___________') else ''
    hd = head if head not in ('', '___________________') else ''
    if ds and hd:
        middle_text = f'{ds}                    {hd}'
    elif ds:
        middle_text = ds
    elif hd:
        middle_text = hd
    else:
        middle_text = ''
    _cell(foot_table.cell(0, 1), middle_text, SZ_MAIN, align=CENTER)
    _set_cell_bottom_border(foot_table.cell(0, 1))

    # Дата — короткий плейсхолдер если не заполнена
    if report_date:
        date_text = f'{_date_signature(report_date)} г.'
    else:
        date_text = '__.__.____  г.'
    _cell(foot_table.cell(0, 2), date_text, SZ_MAIN, align=RIGHT)
    _set_cell_bottom_border(foot_table.cell(0, 2))

    # Строка 1: подсказки под линией — одна строка, шрифт 12pt
    _cell(foot_table.cell(1, 0), '', SZ_HINT)
    _cell(foot_table.cell(1, 1),
          '(подпись, И.О.Фамилия)',
          SZ_HINT, align=CENTER)
    _cell(foot_table.cell(1, 2), '', SZ_HINT)




def generate_individual_plan(position):
    """
    Генерирует Word-документ индивидуального плана для позиции.


    """
    from apps.core.models import ExportSettings
    from apps.tasks.services import (
        calculate_teaching_summary,
        get_position_workload_stats,
    )

    doc = Document()
    _setup(doc)

    # Настройки выгрузки
    try:
        settings = ExportSettings.objects.get(
            academic_year=position.academic_year,
        )
    except ExportSettings.DoesNotExist:
        settings = None

    # Предварительный расчёт
    stats = get_position_workload_stats(position)
    plan_summary = calculate_teaching_summary(position, use_fact=False)
    fact_summary = calculate_teaching_summary(position, use_fact=True)

    plan_total = plan_summary.get('total', Decimal('0'))
    fact_total = fact_summary.get('total', Decimal('0'))

    # ── Сборка страниц ──
    _add_header_page(doc, position, settings)
    _add_summary_page(doc, position, settings, stats,
                      plan_total, fact_total)
    _add_teaching_section(doc, position, plan_summary, fact_summary)

    if not position.is_hourly:
        _add_task_sections(doc, position, stats)

    _add_footer_page(doc, settings)

    # Сохранение
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer